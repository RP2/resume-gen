# resume_generator/utils/pdf.py
"""
PDF generation utilities using WeasyPrint.
"""
from weasyprint import HTML
from docx import Document
from bs4 import BeautifulSoup, Tag
import logging

def html_to_pdf(html: str, output_path: str):
    HTML(string=html).write_pdf(output_path)

def html_to_docx(html: str, output_path: str):
    """Convert HTML resume to Word docx using python-docx."""
    try:
        doc = Document()
        soup = BeautifulSoup(html, "html.parser")
        # Simple conversion: add each block as a paragraph
        for elem in soup.find_all(['h1','h2','h3','p','li','section','header']):
            text = elem.get_text(strip=True)
            if text and isinstance(elem, Tag):
                if elem.name == 'h1':
                    doc.add_heading(text, level=1)
                elif elem.name == 'h2':
                    doc.add_heading(text, level=2)
                elif elem.name == 'h3':
                    doc.add_heading(text, level=3)
                else:
                    doc.add_paragraph(text)
        if not doc.paragraphs:
            doc.add_paragraph("[No resume content found]")
        doc.save(output_path)
    except Exception as e:
        logging.error(f"Error converting HTML to docx: {e}")
        raise

# Add more PDF-related utilities as needed
