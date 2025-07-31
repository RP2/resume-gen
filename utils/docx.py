"""
Word document generation and formatting utilities using python-docx.
"""
from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup, Tag
import logging


def html_to_docx(html: str, output_path: str):
    """Convert HTML resume to Word docx using python-docx, with improved section grouping and deduplication."""
    try:
        doc = Document()
        soup = BeautifulSoup(html, "html.parser")
        seen_sections = set()
        section_map = {"contact": [], "summary": [], "experience": [], "skills": []}
        # Group content by section
        for elem in soup.find_all(['h1','h2','h3','p','li','section','header']):
            text = elem.get_text(strip=True)
            if not text:
                continue
            lower = text.lower()
            if "contact" in lower or "based in" in lower or "@" in text or "http" in text:
                if text not in section_map["contact"]:
                    section_map["contact"].append(text)
            elif "summary" in lower:
                if text not in section_map["summary"]:
                    section_map["summary"].append(text)
            elif "experience" in lower or "developer" in lower or "specialist" in lower:
                if text not in section_map["experience"]:
                    section_map["experience"].append(text)
            elif "skills" in lower or any(skill in lower for skill in ["javascript","python","react","css","html","project management"]):
                if text not in section_map["skills"]:
                    section_map["skills"].append(text)
        # Write sections in order
        if section_map["contact"]:
            doc.add_heading("Contact Information", level=2)
            for line in section_map["contact"]:
                doc.add_paragraph(line)
        if section_map["summary"]:
            doc.add_heading("Summary", level=2)
            for line in section_map["summary"]:
                doc.add_paragraph(line)
        if section_map["experience"]:
            doc.add_heading("Experience", level=2)
            for line in section_map["experience"]:
                doc.add_paragraph(line)
        if section_map["skills"]:
            doc.add_heading("Skills", level=2)
            for line in section_map["skills"]:
                doc.add_paragraph(line)
        # Add extra spacing between sections
        for p in doc.paragraphs:
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.space_before = Pt(3)
        if not doc.paragraphs:
            doc.add_paragraph("[No resume content found]")
        doc.save(output_path)
    except Exception as e:
        logging.error(f"Error converting HTML to docx: {e}")
        raise


def format_docx(doc):
    """Apply custom formatting to a Word Document (fonts, spacing, etc)."""
    # Example: set default font to Arial, size 11
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    # Add more formatting as needed
    return doc

# TODO: DOCX output is experimental and not production-ready. Focus on PDF output for now.
# Improvements needed: section grouping, formatting, deduplication, layout, etc.
# Add more docx-related utilities as needed
